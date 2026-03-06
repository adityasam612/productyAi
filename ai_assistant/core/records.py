# core/records.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from core.tasks import get_tasks

EXPORTS_DIR = "exports"

def export_weekly_report() -> str:
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_heading(f"Weekly Productivity Report — {datetime.now():%B %d, %Y}", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    # Completed tasks
    doc.add_heading("✅ Completed Tasks", level=1)
    completed = get_tasks(done=1)
    if completed:
        for t in completed:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{t[1]}").bold = True
            p.add_run(f" — {t[2]} priority | due: {t[3] or 'none'}")
    else:
        doc.add_paragraph("No tasks completed this period.")

    # Pending tasks
    doc.add_heading("⏳ Pending Tasks", level=1)
    pending = get_tasks(done=0)
    if pending:
        for t in pending:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{t[1]}").bold = True
            p.add_run(f" — {t[2]} priority | due: {t[3] or 'none'}")
    else:
        doc.add_paragraph("No pending tasks.")

    # Stats table
    doc.add_heading("📊 Summary Stats", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Completed"
    table.rows[0].cells[1].text = "Pending"
    table.rows[1].cells[0].text = str(len(completed))
    table.rows[1].cells[1].text = str(len(pending))

    filename = f"{EXPORTS_DIR}/report_{datetime.now():%Y_%m_%d_%H%M}.docx"
    doc.save(filename)
    return filename